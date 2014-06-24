#!/usr/bin/env python
# -*- coding: utf-8 -*-
'''
Copyright (C) 2014  v4lproik

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.
'''


'''
This program aims to parse nmap xml files.
It creates a .docx with a table which contains the different hosts scanned and their services.
'''

__author__ = "v4lproik"
__date__ = "24/06/2014"
__version__ = "1.0"
__maintainer__ = "v4lproik"
__email__ = "v4lproik@gmail.com"
__status__ = "Development"
__twitter__ = "v4lproik"


try:
	import traceback
	import sys
	import nmap
	import argparse
	from docx import Document
	from docx.shared import Inches
	import os
	import glob
	import re
except:
		print traceback.format_exc()

DEBUG = False

def main(file):

	#structure to pass to create doc
	data = [["","",[],[],[],[],[],[]]]

	for f in file:
		nm = nmap.PortScanner()
		nm.analyse_nmap_xml_scan(open(f).read())
		#print nm.scaninfo()


		for host in nm.all_hosts():

			hostO = ""
			os = ""
			ports = []
			services = []
			softs = []
			versions = []
			information = []

			#print nm[host]
			oHost = nm[host]

			index = 0
			for infohost in data:
				#print oHost

				if infohost[0].startswith(host):
					ports = infohost[2]
					services = infohost[3]
					softs = infohost[4]
					versions = infohost[5]
					information = infohost[6]
					data.remove(infohost)
					break;
				index+=1


			if oHost.state() == "up":
				#print oHost.state() + " => " + host + "(" + oHost.hostname() + ")"
				hostO = host + "(" + oHost.hostname() + ")"

				try:
					os = nm[host]['osmatch'][0]['name']
				except KeyError:
					pass

				if(len(nm[host].all_protocols()))>1 or np:
					for proto in nm[host].all_protocols():
						if proto == "tcp" or proto == "udp":
							lport = nm[host][proto].keys()
							lport.sort()

							for port in lport:
								#print nm[host][proto][port]
								#ports.append(str(port) + " " + str(nm[host][proto][port]['state']) + " " + str(proto))
								portL = str(port) + " " + str(nm[host][proto][port]['state']) + " " + str(proto)
								service = str(nm[host][proto][port]['name'])
								soft = str(nm[host][proto][port]['product'])
								version = str(nm[host][proto][port]['version'])
								extrainfo = str(nm[host][proto][port]['extrainfo'])

								ports, services, softs, versions, information = sort_ports(ports, services, softs, versions, information, portL, service, soft, version, extrainfo)
								#services.append(str(nm[host][proto][port]['name']))
								#softs.append(str(nm[host][proto][port]['product']))
								#versions.append(str(nm[host][proto][port]['version']))

					data.append([hostO, os, ports, services, softs, versions, information])

	return data


def sort_ports(portsL, services, softs, versions, infos, port, service, soft, version, info):
	tmp_ports = []
	tmp_services = []
	tmp_softs = []
	tmp_versions = []
	tmp_infos = []

	count = 0
	if len(portsL) > 0:
		flag = False
		for ports in portsL:
			p = ports.split(" ")[0]
			if int(port.split(" ")[0]) < int(p) and not flag:
				tmp_ports.append(port)
				tmp_services.append(service)
				tmp_softs.append(soft)
				tmp_versions.append(version)
				tmp_infos.append(info)
				flag = True

			#print ports
			tmp_ports.append(portsL[count])
			tmp_services.append(services[count])
			tmp_softs.append(softs[count])
			tmp_versions.append(versions[count])
			tmp_infos.append(infos[count])
			count+=1

		if not flag:
			tmp_ports.append(port)
			tmp_services.append(service)
			tmp_softs.append(soft)
			tmp_versions.append(version)
			tmp_infos.append(info)

	else:
		tmp_ports.append(port)
		tmp_services.append(service)
		tmp_softs.append(soft)
		tmp_versions.append(version)
		tmp_infos.append(info)

	return tmp_ports,tmp_services,tmp_softs,tmp_versions,tmp_infos


def report():
	document = Document()
	table = document.add_table(rows=1, cols=7)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Hote'
	hdr_cells[1].text = 'Systeme dexploitation'
	hdr_cells[2].text = 'Port'
	hdr_cells[3].text = 'Service'
	hdr_cells[4].text = 'Logiciel'
	hdr_cells[5].text = 'Version'
	hdr_cells[6].text = 'Information'


	table2 = document.add_table(rows=1, cols=1)
	#print data
	count = 0
	for i in range(0,len(data),1):
		line = data.pop()
		#print line
		row_cells = table.add_row().cells
		row_cells[0].text = line[0]
		row_cells[1].text = line[1]

		count2 = 0
		for i in range(len(line[2])):

			if count2 > 0:
				row_cells = table.add_row().cells
				row_cells[0].text = ""
				row_cells[1].text = ""
			row_cells[2].text = line[2][count2]
			row_cells[3].text = line[3][count2]
			row_cells[4].text = line[4][count2]
			row_cells[5].text = line[5][count2]
			row_cells[6].text = line[6][count2]
			count2+=1

	if tablestyle:
		table.style=tablestyle

	document.add_page_break()

	document.save(o)
	print "[*] Jobs done\n"


def banner():
	banner = '''
	|----------------------------------------------------------|
	|              Parser Nmap Output to Docx 1.0              |
	|                         V4lproik                         |
	|----------------------------------------------------------|\n'''
	print banner


if __name__ == '__main__':
	#try:
	parser = argparse.ArgumentParser()
	gr1 = parser.add_argument_group("main arguments")
	gr1.add_argument('-f', '--filename', dest='filename',
					 required=True, help='Nmap output file or folder containing several .xml. If its a folder all .xml will be taking into accoutn.')
	gr1.add_argument('-o', '--output', dest='output',
					 required=True, help='docx file location')

	gr2 = parser.add_argument_group("optional arguments")
	gr2.add_argument('-np', '--no-port', dest='noport',
					 default=False,  action='store_true', help='Include hosts that are up with no ports opened found')
	gr2.add_argument('-s', '--table-style', dest='tablestyle',
					 default=False, help='Specify your word table style')

	banner()

	args = parser.parse_args()

	fi = args.filename
	o = args.output
	np = args.noport
	tablestyle = args.tablestyle
	folder = False
	if os.path.isdir(fi):
		f = glob.glob(fi + "/*.xml")
		if len(f) < 1:
			sys.exit("No nmap xml files have been found within the following directory : " + fi)
		else:
			print "[*] " + str(len(f)) + " nmap xml files have been found."
	else:
		f=fi

	data = main(f)
	report()


	sys.exit(0)
	#except:
		#print traceback.format_exc()